from fastapi import APIRouter, HTTPException
from app.services.relatorio import Gerar_relatorio_deepseek, Gerar_relatorio_gemini, Gerar_relatorio_groq, Gerar_relatorio_openai
from typing import Any
from docx import Document
from fastapi.responses import FileResponse
import os

relatorioRouter = APIRouter(prefix="/relatorio")

questoes = [
    "Com saber se é Alzheimer, estresse ou processo natural do envelhecimento?",
    "Quais exames detectam Alzheimer?",
    "Que especialidade médica pode solicitar exames e dar diagnóstico?",
    "O que vai acontecer, ele vai esquecer - só isso?",
    "Tem cura? Tem tratamento para diminuir os sintomas ou retardar o avanço da doença?",
    "Ele tem quanto tempo de vida?",
    "É hereditário?",
    "Ele está estranho, mas continua trabalhando e dirigindo, será que é Alzheimer mesmo?",
    "Ele não quer parar de dirigir, o que eu faço? Ele quase atropelou uma senhora quando subiu na calçada.",
    "Está ficando agressivo e nega que tem alguma doença. Como impedir que saia sozinho?",
    "Não quer ir ao médico, o que eu faço? Como posso convencê-lo?",
    "A agressividade ou a sexualidade exacerbada tem algum medicamento para controlar?",
    "A agitação e os delírios são normais? Ele vê fantasmas e bichos na parede, como devo lidar com isso, sem deixá-lo agressivo?",
    "Ele acaba de almoçar e quer almoçar novamente, diz que não comeu nada, ou ele se recusa a comer qualquer alimento e beber água. Existe alguma vitamina ou produto que supra a falta de alimento?",
    "Ele não quer tomar banho, como fazer a higiene dele, se não tenho forças para levá-lo ao chuveiro à força?",
    "Ele me perguntou quem sou eu. Não sabe que é meu pai, em algum momento ele vai me reconhecer novamente?",
    "Quer voltar para casa, insiste que aqui não é a casa dele, como acalmá-lo e mostrar que está na sua casa?",
    "Final da tarde ele começa a ficar irritado, como prevenir ou diminuir esse estresse e medo do entardecer?",
    "Não consegue mais escovar os dentes, como faço a higiene bucal dele?",
    "Ele não consegue mais se alimentar, precisa colocar sonda, como vou manusear essa sonda?",
    "Ele está ficando com o corpo rígido e anda se inclinando, caiu várias vezes, existe algum tratamento para ajudar o equilíbrio?",
    "Machuquei o paciente quando tirei da cama e coloquei na cadeira de rodas, desloquei meu ombro nessa manobra. Existe alguma técnica para mover o paciente sem machucar ele e o cuidador?",
    "Ele só dorme, não consegue ficar acordado e passa o dia e noite dormindo, isso é normal da doença?",
    "Ele não dorme, delira o tempo todo e está esgotado, que tipo de medicamento pode ajudar a acalmar e dormir?",
    "Está depressivo e só fala em morrer, adiantaria procurar um terapeuta?",
    "Não fala mais, está silencioso, olhar parado e parece não reagir à minha presença, o que eu faço?",
    "Pneumonia, Infecção urinária, gripe forte... ele não tem resistência e o corpo está falindo. Como eu posso salvá-lo?",
    "Ele está muito mal, mas eu não quero levá-lo ao hospital porque quero perto da família. E se ele morrer em casa, serei culpada pela morte dele? Se ele morrer no hospital longe da família, serei culpada pelo sofrimento dele?",
    "Estou depressiva e esgotada, gostaria de morrer no lugar dele. Onde consigo ajuda para sair deste poço escuro ?",
    "Tenho saudades, será que poderia ter feito melhor?",
    "Acordo assustada ouvindo-o me chamar, mesmo depois de 2 anos que faleceu, preciso de ajuda. Onde consigo profissional que me ajude?",
    "Ele foi embora e a casa está vazia, minha vida está vazia. Como consigo recomeçar com esse peso?",
    "Esqueci o que fui buscar no mercado, será que estou no início do Alzheimer?",
]

template_path = "app/template/template-relatorio.docx"
output = "app/output/"

@relatorioRouter.get("/deepseek")
async def chat() -> Any:
    relatorio = output+"relatorio-deepseek.docx"
    if not os.path.exists(relatorio):
        await GerarRelatorio(nome_llm="Deepseek", output=relatorio)
        
    return FileResponse(
        path=relatorio,
        filename="relatorio_deepseek.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
@relatorioRouter.get("/gemini")
async def chat() -> Any:
    relatorio = output+"relatorio-gemini.docx"
    if not os.path.exists(relatorio):
        await GerarRelatorio(nome_llm="Gemini", output=relatorio)
    return FileResponse(
        path=relatorio,
        filename="relatorio_gemini.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
@relatorioRouter.get("/openai")
async def chat() -> Any:
    relatorio = output+"relatorio-openai.docx"
    
    if not os.path.exists(relatorio):
        await GerarRelatorio(nome_llm="Openai", output=relatorio)
    return FileResponse(
        path=relatorio,
        filename="relatorio_openai.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
@relatorioRouter.get("/groq")
async def chat() -> Any:
    relatorio = output+"relatorio-groq.docx"
    if not os.path.exists(relatorio):    
        await GerarRelatorio(nome_llm="Groq", output=relatorio)
    return FileResponse(
        path=relatorio,
        filename="relatorio_groq.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def substituir_campos(template_path, output_path, substituicoes):
    try:
        doc = Document(template_path)
        
        for paragraph in doc.paragraphs:
            for campo, valor in substituicoes.items():
                if campo in paragraph.text:
                    print(f"Chave {campo}, valor: {valor}\n")
                    paragraph.text = paragraph.text.replace(campo, str(valor))
        
        # Substitui texto em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for campo, valor in substituicoes.items():
                        if campo in cell.text:                    
                            print(f"Chave {campo}, valor: {valor}\n")
                            cell.text = cell.text.replace(campo, str(valor))
        doc.save(output_path)
        print(f"Documento gerado com sucesso em: {output_path}")
    except Exception as e:
        print(f"Erro ao processar o documento: {str(e)}")
        
async def GerarRelatorio(nome_llm, output):
    substituir_campos(template_path=template_path, output_path=output, substituicoes={"$$LLM$$": nome_llm})
    resultados = []
    for questao in questoes:
        try:
            if nome_llm == "Deepseek":
                response = await Gerar_relatorio_deepseek(questao)
                resultados.append(response)
            if nome_llm == "Gemini":
                response = await Gerar_relatorio_gemini(questao)
                resultados.append(response)
            if nome_llm == "Openai":
                response = await Gerar_relatorio_openai(questao)
                resultados.append(response)
            if nome_llm == "Groq":
                response = await Gerar_relatorio_groq(questao)
                resultados.append(response)
                    
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Erro interno: {str(e)}")

    for i in range(len(resultados)):
        resposta = resultados[i]
        substituir_campos(template_path=output, output_path=output,
                        substituicoes={
                            f"$$pergunta{i}$$": resposta["pergunta"],
                            f"$$resposta_com_rag_{i}$$":resposta["resposta_com_rag"],
                            f"$$resposta_sem_rag_{i}$$": resposta["resposta_sem_rag"]})
